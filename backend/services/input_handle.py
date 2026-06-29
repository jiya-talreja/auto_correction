from pdf2image import convert_from_bytes
from services.img_ext import run_ocr
from docx import Document
import io
from PIL import Image
import fitz
POPPLER_PATH = r"C:\Users\DELL\Downloads\Release-25.12.0-0\poppler-25.12.0\Library\bin"
def pdf(file_bytes):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        page_text = page.get_text().strip()
        if page_text:
            text += page_text + "\n"
    if text.strip():
        print("TEXT PDF")
        return text
    print("SCANNED PDF")
    images = convert_from_bytes(
        file_bytes,
        dpi=300,
        poppler_path=POPPLER_PATH
    )
    ocr_text = ""
    for img in images:
        page_text = run_ocr(img)      # run_ocr should also return a STRING
        ocr_text += page_text + "\n"
    return ocr_text
def docs(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    return text
def img_in(img_bytes_list):
    text = ""
    for img_b in img_bytes_list:
        img = Image.open(io.BytesIO(img_b))
        text += run_ocr(img)
        text += "\n"
    return text
  
